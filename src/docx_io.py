"""DOCX extraction and minimal span-level editing helpers."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from .models import AcceptedDetection, TextLocation


@dataclass
class LocationBinding:
    location: TextLocation
    runs: list[object] | None = None
    text_nodes: list[object] | None = None


def load_document(path: Path) -> DocumentType:
    return Document(str(path))


def _location(key: str, kind: str, text: str, **coordinates: object) -> TextLocation:
    context = coordinates.pop("context", ())
    return TextLocation(key=key, kind=kind, text=text, context=tuple(context), **coordinates)


def _append_paragraph(bindings: dict[str, LocationBinding], key: str, doc_paragraph: Paragraph, kind: str, **coordinates: object) -> None:
    text = "".join(run.text for run in doc_paragraph.runs)
    if text:
        loc = _location(key, kind, text, **coordinates)
        bindings[key] = LocationBinding(loc, runs=list(doc_paragraph.runs))


def _walk_table(bindings: dict[str, LocationBinding], table: Table, table_number: int, section: int | None, prefix: str, seen_cells: list[object]) -> None:
    headers = [" ".join(cell.text.split()) for cell in table.rows[0].cells] if table.rows else []
    for row_index, row in enumerate(table.rows, start=1):
        for cell_index, cell in enumerate(row.cells, start=1):
            if any(existing is cell._tc for existing in seen_cells):
                continue
            seen_cells.append(cell._tc)
            header = headers[cell_index - 1] if cell_index <= len(headers) else ""
            context = (("table_headers", " | ".join(headers)), ("column_header", header))
            for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                _append_paragraph(bindings, f"{prefix}.r{row_index}.c{cell_index}.p{paragraph_index}", paragraph, "table_cell", section=section, table=table_number, row=row_index, cell=cell_index, paragraph=paragraph_index, context=context)
            for nested_index, nested in enumerate(cell.tables, start=1):
                _walk_table(bindings, nested, table_number, section, f"{prefix}.nested{nested_index}", seen_cells)


def extract_locations(document: DocumentType) -> tuple[dict[str, LocationBinding], list[str]]:
    bindings: dict[str, LocationBinding] = {}
    warnings: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        _append_paragraph(bindings, f"body.p{index}", paragraph, "paragraph", paragraph=index)
    seen_cells: list[object] = []
    for table_index, table in enumerate(document.tables, start=1):
        _walk_table(bindings, table, table_index, None, f"body.t{table_index}", seen_cells)

    seen_parts: set[str] = set()
    for section_index, section in enumerate(document.sections, start=1):
        for kind, container in (("header", section.header), ("footer", section.footer)):
            part_name = str(container.part.partname)
            if part_name in seen_parts:
                continue
            seen_parts.add(part_name)
            for paragraph_index, paragraph in enumerate(container.paragraphs, start=1):
                _append_paragraph(bindings, f"{kind}.{part_name}.p{paragraph_index}", paragraph, kind, section=section_index, paragraph=paragraph_index, part=part_name)
            for table_index, table in enumerate(container.tables, start=1):
                _walk_table(bindings, table, table_index, section_index, f"{kind}.{part_name}.t{table_index}", seen_cells)
            _extract_textboxes(bindings, container.part._element, part_name, kind, section_index)
    _extract_textboxes(bindings, document.part._element, str(document.part.partname), "textbox", None)

    image_parts = [part for part in document.part.package.parts if "image" in str(getattr(part, "content_type", ""))]
    if image_parts:
        warnings.append(f"Image content exists ({len(image_parts)} image parts); OCR is not supported.")
    return bindings, warnings


def _extract_textboxes(bindings: dict[str, LocationBinding], element: object, part_name: str, kind: str, section: int | None) -> None:
    try:
        boxes = element.xpath(".//*[local-name()='txbxContent']")
    except AttributeError:
        return
    for index, box in enumerate(boxes, start=1):
        nodes = list(box.xpath(".//*[local-name()='t']"))
        text = "".join(node.text or "" for node in nodes)
        if text:
            key = f"textbox.{part_name}.{index}"
            loc = _location(key, "textbox", text, section=section, part=part_name, context=(("source", kind),))
            bindings[key] = LocationBinding(loc, text_nodes=nodes)


def apply_replacements(bindings: dict[str, LocationBinding], replacements: Iterable[tuple[AcceptedDetection, str]]) -> list[str]:
    warnings: list[str] = []
    grouped: dict[str, list[tuple[AcceptedDetection, str]]] = {}
    for detection, value in replacements:
        grouped.setdefault(detection.candidate.location.key, []).append((detection, value))
    for key, items in grouped.items():
        binding = bindings.get(key)
        if binding is None:
            warnings.append(f"Missing location binding for {key}.")
            continue
        pieces = binding.runs if binding.runs is not None else binding.text_nodes
        if not pieces:
            warnings.append(f"Unsupported empty binding for {key}.")
            continue
        for detection, replacement in sorted(items, key=lambda item: item[0].candidate.start, reverse=True):
            if not _replace_in_pieces(pieces, detection.candidate.start, detection.candidate.end, replacement):
                warnings.append(f"Could not safely apply {detection.detection_id} at {key}.")
    return warnings


def _piece_text(piece: object) -> str:
    return getattr(piece, "text", "") or ""


def _set_piece_text(piece: object, value: str) -> None:
    setattr(piece, "text", value)


def _replace_in_pieces(pieces: list[object], start: int, end: int, replacement: str) -> bool:
    positions: list[tuple[int, int]] = []
    cursor = 0
    for piece in pieces:
        next_cursor = cursor + len(_piece_text(piece))
        positions.append((cursor, next_cursor))
        cursor = next_cursor
    if start < 0 or end > cursor or start >= end:
        return False
    affected = [index for index, (left, right) in enumerate(positions) if start < right and left < end]
    if not affected:
        return False
    first, last = affected[0], affected[-1]
    for index in range(first, last + 1):
        left, right = positions[index]
        value = _piece_text(pieces[index])
        before = value[: max(0, start - left)] if index == first else ""
        after = value[max(0, end - left) :] if index == last else ""
        _set_piece_text(pieces[index], before + (replacement if index == first else "") + after)
    return True


def save_and_reopen(document: DocumentType, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    Document(str(path))
