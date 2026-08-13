from __future__ import annotations

from pathlib import Path
import json
import tempfile
import unittest

from docx import Document

from src.redactor import build_parser, run


class RedactorIntegrationTests(unittest.TestCase):
    def test_cli_creates_output_and_audit_without_redacting_financial_year(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "input.docx"
            output = root / "out.docx"
            audit = root / "audit.json"
            document = Document()
            document.add_paragraph("Contact Person: John Smith; Telephone: +91 90000 00001; Email: john.smith@example.com")
            document.add_paragraph("Fiscal 2022-2023 revenue was ₹ 1,000 million.")
            document.save(source)
            args = build_parser().parse_args([str(source), "--output", str(output), "--audit", str(audit)])
            payload, code = run(args)
            self.assertEqual(0, code)
            self.assertTrue(output.exists())
            self.assertTrue(audit.exists())
            result = Document(output)
            text = "\n".join(paragraph.text for paragraph in result.paragraphs)
            self.assertNotIn("john.smith@example.com", text)
            self.assertNotIn("John Smith", text)
            self.assertIn("2022-2023", text)
            self.assertGreaterEqual(payload["counts_by_category"].get("EMAIL", 0), 1)
            audit_data = json.loads(audit.read_text(encoding="utf-8"))
            self.assertNotIn("john.smith@example.com", audit.read_text(encoding="utf-8"))
            self.assertIn("detected", audit_data)
