from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document

from src.docx_io import apply_replacements, extract_locations, save_and_reopen
from src.models import AcceptedDetection, CandidateSpan


class DocxIoTests(unittest.TestCase):
    def test_extracts_and_replaces_cross_run_paragraph_table_header_footer_and_nested_table(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Contact Person: John ").bold = True
        paragraph.add_run("Smith")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "jane@example.com"
        nested = table.cell(0, 0).add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "Telephone: +91 90000 00001"
        document.sections[0].header.paragraphs[0].text = "Email: header@example.com"
        document.sections[0].footer.paragraphs[0].text = "Email: footer@example.com"
        bindings, _ = extract_locations(document)
        target = next(binding for binding in bindings.values() if binding.location.text == "Contact Person: John Smith")
        start = target.location.text.index("John Smith")
        detection = AcceptedDetection("det-0001", CandidateSpan("PERSON_NAME", target.location, start, start + len("John Smith"), "John Smith", "john smith", "high", ("test",)))
        warnings = apply_replacements(bindings, [(detection, "Aarav Mehta")])
        self.assertEqual([], warnings)
        self.assertEqual("Contact Person: Aarav Mehta", "".join(run.text for run in paragraph.runs))
        self.assertTrue(paragraph.runs[0].bold)
        self.assertTrue(any(binding.location.kind == "table_cell" for binding in bindings.values()))
        self.assertTrue(any(binding.location.kind == "header" for binding in bindings.values()))
        self.assertTrue(any(binding.location.kind == "footer" for binding in bindings.values()))
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "result.docx"
            save_and_reopen(document, output)
            self.assertTrue(output.exists())


class DocxIoMergedCellTests(unittest.TestCase):
    """Regression tests for merged-table-cell traversal deduplication.

    In python-docx, all logical positions of a merged cell share the same
    underlying _tc lxml element.  The _walk_table implementation uses
    ``cell._tc`` identity (``is``) to skip already-visited cells.  These tests
    verify that a merged cell is traversed exactly once, that PII inside it is
    detected exactly once, and that replacement succeeds without duplication.
    """

    def _build_merged_table_document(self) -> Document:
        """Return an in-memory Document with a 3-column table where columns 0
        and 1 are merged into one physical cell containing an email address."""
        document = Document()
        table = document.add_table(rows=2, cols=3)
        # Row 0: header row
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Name"   # will be the same _tc after merge
        table.cell(0, 2).text = "Contact"
        # Merge columns 0 and 1 in row 1, place PII in the merged cell
        merged = table.cell(1, 0).merge(table.cell(1, 1))
        merged.text = "E-mail: merged.cell@example.com"
        table.cell(1, 2).text = "Unrelated text"
        return document

    def test_merged_cell_traversed_exactly_once(self) -> None:
        """A merged cell must produce one LocationBinding, not one per logical position."""
        document = self._build_merged_table_document()
        bindings, _ = extract_locations(document)
        # Find all bindings whose text contains the merged-cell email
        email_bindings = [
            key for key, binding in bindings.items()
            if "merged.cell@example.com" in binding.location.text
        ]
        self.assertEqual(
            1, len(email_bindings),
            f"Expected exactly 1 binding for the merged cell, got {len(email_bindings)}: {email_bindings}",
        )

    def test_merged_cell_pii_detected_exactly_once(self) -> None:
        """Email PII inside a merged cell must be detected exactly once (no duplicates)."""
        from src.detectors import detect_locations, select_candidates

        document = self._build_merged_table_document()
        bindings, _ = extract_locations(document)
        accepted, _ = select_candidates(detect_locations(b.location for b in bindings.values()))
        email_detections = [
            d for d in accepted
            if d.candidate.category == "EMAIL"
            and "merged.cell@example.com" in d.candidate.normalized_value
        ]
        self.assertEqual(
            1, len(email_detections),
            f"Expected exactly 1 EMAIL detection for merged cell, got {len(email_detections)}",
        )

    def test_merged_cell_replacement_succeeds_and_docx_reopens(self) -> None:
        """Replacement in a merged cell must apply cleanly and the DOCX must reopen."""
        from src.detectors import detect_locations, select_candidates

        document = self._build_merged_table_document()
        bindings, _ = extract_locations(document)
        accepted, _ = select_candidates(detect_locations(b.location for b in bindings.values()))
        email_detections = [
            d for d in accepted
            if d.candidate.category == "EMAIL"
            and "merged.cell@example.com" in d.candidate.normalized_value
        ]
        self.assertEqual(1, len(email_detections))
        warnings = apply_replacements(bindings, [(email_detections[0], "user001@example.test")])
        self.assertEqual([], warnings, f"Unexpected replacement warnings: {warnings}")
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "merged_result.docx"
            save_and_reopen(document, output)
            self.assertTrue(output.exists())
            # Verify original email is gone and replacement is present
            reopened = Document(str(output))
            all_text = "\n".join(
                cell.text for table in reopened.tables for row in table.rows for cell in row.cells
            )
            self.assertNotIn("merged.cell@example.com", all_text)
            self.assertIn("user001@example.test", all_text)

