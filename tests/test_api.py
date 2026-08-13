import io
import unittest
from fastapi.testclient import TestClient
from docx import Document

from src.api import app

client = TestClient(app)

class ApiTests(unittest.TestCase):
    def test_health_check(self) -> None:
        response = client.get("/health")
        self.assertEqual(200, response.status_code)
        self.assertEqual({"status": "ok"}, response.json())
        
    def test_redact_requires_docx_extension(self) -> None:
        response = client.post("/redact", files={"file": ("test.txt", b"hello world")})
        self.assertEqual(400, response.status_code)
        self.assertIn("Only .docx files", response.json()["detail"])
        
    def test_redact_requires_non_empty_file(self) -> None:
        response = client.post("/redact", files={"file": ("test.docx", b"")})
        self.assertEqual(400, response.status_code)
        self.assertIn("empty", response.json()["detail"])
        
    def test_redact_processes_docx_successfully(self) -> None:
        # Create a small valid DOCX in memory
        doc = Document()
        doc.add_paragraph("Contact: john.doe@example.com")
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_bytes = doc_io.getvalue()
        
        response = client.post("/redact", files={"file": ("test.docx", doc_bytes)})
        self.assertEqual(200, response.status_code)
        self.assertEqual("application/vnd.openxmlformats-officedocument.wordprocessingml.document", response.headers["content-type"])
        self.assertTrue(response.headers["content-disposition"].startswith('attachment; filename="redacted_test.docx"'))
        
        # Verify the returned file is a valid DOCX and redacted
        result_io = io.BytesIO(response.content)
        result_doc = Document(result_io)
        text = "\n".join(p.text for p in result_doc.paragraphs)
        self.assertNotIn("john.doe@example.com", text)
        self.assertIn("@example.test", text)
